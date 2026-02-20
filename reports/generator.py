# reports/generator.py

import os
from docx import Document
from rag.generator import generate_answer
from datetime import datetime
import tempfile

# Standard analyst questions
ANALYST_QUESTIONS = {
    "Company Overview": "Summarize the company’s business model and main revenue sources.",
    "Financial Performance": "Analyze revenue, profitability, and margin trends over recent years.",
    "Growth Outlook": "Analyze the company’s growth drivers and future prospects.",
    "Key Risks": "Identify major business, financial, and operational risks.",
    "Management Outlook": "Summarize management strategy and future guidance if available.",
    "Competitive Position": "Analyze the company’s position relative to competitors.",
    "Investment Thesis": "Provide an overall investment thesis based on the report."
}


def generate_report(company_name: str, output_path: str = None):
    """
    Generates an investment memo using RAG system
    and saves it as a DOCX file.
    """

    print("Starting report generation...\n")

    sections = {}

    # Ask questions automatically
    for section, question in ANALYST_QUESTIONS.items():
        print(f"Generating section: {section}")
        full_query = f"{question} Focus on {company_name}."
        answer = generate_answer(full_query)
        sections[section] = answer

    # Create document
    doc = Document()

    # Title
    doc.add_heading(
        f"Investment Memo — {company_name}",
        level=1
    )

    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    doc.add_paragraph("-" * 40)

    # Add sections
    for section, content in sections.items():
        doc.add_heading(section, level=2)
        for line in content.split("\n"):
            doc.add_paragraph(line)

    # Use temp file if output_path not provided (works on Streamlit Cloud)
    if output_path is None:
        tmp_file = tempfile.NamedTemporaryFile(suffix="_memo.docx", delete=False)
        output_path = tmp_file.name

    doc.save(output_path)
    print(f"\nReport saved at: {output_path}")

    return output_path